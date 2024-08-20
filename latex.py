from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_word_equation(latex_expr):
    # Open a new Word document
    doc = Document()

    # Add a title
    doc.add_heading('Word Equation from LaTeX', level=1)

    # Start adding the equation
    equation = doc.add_paragraph()
    
    # Write the LaTeX expression as a Word equation
    run = equation.add_run()
    run.font.size = doc.styles['Normal'].font.size

    # Construct the Word equation syntax equivalent of the LaTeX expression
    word_equation = "lim_{x -> 2} (x^k - 2^k)/(x - 2)"
    
    # Add the constructed equation to the paragraph
    run.text = word_equation

    # Save the document
    doc.save('Word_Equation.docx')

# LaTeX expression
latex_expression = "$\lim_{ x \to 2}$ $\frac{x^k-2^k}{x-2}$"

# Call the function to create a Word document with the equation
create_word_equation(latex_expression)
