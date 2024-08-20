import docx
import json

file = open("q.json", encoding="utf8")
data = json.load(file)
print(data["questions"][0]["question"])
file.close()

# Create a document
doc = docx.Document()

# Add a paragraph to the document
p = doc.add_paragraph()

# Add some formatting to the paragraph
p.paragraph_format.line_spacing = 1
p.paragraph_format.space_after = 0

# Add a run to the paragraph
run = p.add_run("python-docx")

# Add some formatting to the run
run.bold = True
run.italic = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)

# Add more text to the same paragraph
run = p.add_run(" Tutorial")

# Format the run
run.bold = True
run.font.name = 'Arial'
run.font.size = docx.shared.Pt(16)
for question in data["questions"]:

  # Add another paragraph
  p = doc.add_paragraph()
  answer = question["answer"]
  explaination = question["explanation"]; 
  if explaination:
    # Add a run and format it
    run = p.add_run(f'''
    {question["sn"]}) {chr(answer+97)} 
    {explaination}
                    ''')
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
  else:
    # Add a run and format it
    run = p.add_run(f'''{question["sn"]}) {chr(answer+97)}''')
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)


# Save the document
doc.save("answer.docx")
